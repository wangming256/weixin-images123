#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# 标题
title = doc.add_heading('民事起诉状', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# 原被告信息
p = doc.add_paragraph()
p.add_run('原告：').bold = True
p.add_run('[你的姓名]，性别：[男/女]，[出生年月日]，民族：[XX]，身份证号：[XXXXXXXXXXXXXXXXXX]，住址：[XX省XX市XX区XX小区XX号楼XX单元XX室]，联系电话：[XXXXXXXXXXX]')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('被告：').bold = True
p.add_run('[物业公司全称]，统一社会信用代码：[XXXXXXXXXXXXXXXXXX]，住所地：[XX省XX市XX区XX小区XX号楼XX单元XX室]，法定代表人：[XXX]，职务：[总经理/董事长]，联系电话：[XXXXXXXXXXX]')

doc.add_paragraph()
doc.add_heading('诉讼请求', level=2)

doc.add_paragraph('1.  请求依法确认被告以未缴纳物业费为由拒绝为原告电梯卡每月授权升级的行为无效；')
doc.add_paragraph('2.  请求判令被告立即为原告电梯卡办理正常授权升级，保障原告正常使用电梯；')
doc.add_paragraph('3.  请求判令被告赔偿原告经济损失人民币[XXXX]元（如有，比如因无法正常使用电梯产生的出行费用、误工费等）；')
doc.add_paragraph('4.  本案诉讼费用由被告承担。')

doc.add_paragraph()
doc.add_heading('事实与理由', level=2)

# 一、基本案情
p = doc.add_paragraph()
p.add_run('一、基本案情').bold = True

doc.add_paragraph(
'\n原告系[XX市XX区XX小区]XX号楼XX单元XX室业主，依法对房屋享有所有权，并有权正常使用小区公共设施包括电梯。\n\n'
'该小区电梯卡采用**每月授权升级制度**：业主物业费缴至当月，物业才给电梯卡授权升级到当月；物业费未缴至当月，物业拒绝授权，电梯卡到期后无法使用电梯。\n\n'
'原告[简述情况：因故暂缓缴纳本期物业费/因物业服务瑕疵未全额缴纳等]，被告据此拒绝为原告电梯卡进行每月授权升级，导致原告电梯卡到期后无法使用电梯，'
'原告**不得不每月多次爬楼梯，严重影响正常出行**，多次沟通被告仍坚持不缴费就不授权，原告无奈只得诉诸法院。'
)

# 二、被告行为违法
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('二、被告行为违反法律规定，构成侵权').bold = True

doc.add_paragraph()
doc.add_paragraph('1. 使用电梯是业主法定基本权利，与物业费缴纳相互独立')
doc.add_paragraph(
'• 业主对小区公共设施（包括电梯）享有共有和共同管理的权利，该项权利由《民法典》直接赋予，'
'基于建筑物区分所有权产生，**不会因为业主欠缴物业费而被剥夺**。\n'
'• 缴纳物业费是业主基于物业服务合同承担的合同义务，与电梯使用权是**两种完全独立的法律关系**。\n'
'• 被告无权将"提供电梯正常使用"作为催缴物业费的筹码，更无权通过"每月不升级就不能用"的方式变相胁迫业主缴费。'
)

doc.add_paragraph()
doc.add_paragraph('2. 该行为本质就是限制使用公共设施，违反民法典明确禁止性规定')
doc.add_paragraph(
'《中华人民共和国民法典》第九百四十四条第三款明确规定：\n\n'
'> **业主逾期不支付物业费的，物业服务人可以催告其在合理期限内支付；合理期限届满仍不支付的，可以提起诉讼或者申请仲裁。**\n'
'> **物业服务人不得采取停止供电、供水、供热、供燃气等方式催交物业费。**\n\n'
'电梯是高层住宅业主日常生活必不可少的公共设施，"拒绝每月授权升级→到期不能用电梯"，'
'这种做法在本质上和停止供水供电没有区别，都是通过停止供应业主必需的公共服务来逼迫业主缴费，'
'完全符合该条款禁止的行为类型，明显违反法律强制性规定。'
)

doc.add_paragraph()
doc.add_paragraph('3. 持续侵权，严重影响原告正常生活')
doc.add_paragraph(
'原告居住在[X]层，被告这种"每月卡脖子"的做法，导致原告每月都面临要么被迫缴费要么爬楼梯的两难选择，'
'严重影响原告日常生活，[如有特殊情况可注明："家中有老人/孕妇/病人，根本无法长期爬楼梯"]，'
'已经对原告正常居住权益造成实质性损害，构成侵权。'
)

# 三、原告损失
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('三、原告损失').bold = True
doc.add_paragraph(
'因被告拒绝授权，原告不得不[简述损失：如"偶尔外出只能打车产生额外交通费每月XXX元"/"爬楼梯导致关节损伤花费医药费XXX元"]，'
'截至起诉之日共计损失人民币[XXXX]元，依法应由被告承担赔偿责任。'
)

doc.add_paragraph()
doc.add_heading('证据清单', level=2)

doc.add_paragraph('1.  业主身份证明：房产证/购房合同/不动产权证——证明原告系小区业主，主体适格；')
doc.add_paragraph('2.  电梯卡需要每月升级、被告拒绝升级的证据：小区电梯卡规则公告/与被告沟通记录（微信/短信/录音/录像）——证明被告实施了拒绝授权的行为；')
doc.add_paragraph('3.  [损失证据：交通费票据/医药费票据等——证明原告实际损失]；')
doc.add_paragraph('4.  其他：[如有其他证据补充]。')

doc.add_paragraph()
doc.add_heading('法律依据', level=2)
doc.add_paragraph('- 《中华人民共和国民法典》第二百七十一条、第九百四十四条、第一千一百六十五条；')
doc.add_paragraph('- 《物业管理条例》第六条。')

doc.add_paragraph()

doc.add_paragraph('综上所述，被告以未缴纳物业费为由拒绝为原告电梯卡办理每月授权升级，违反民法典明确禁止性规定，严重侵害原告正常居住权益，恳请贵院依法支持原告诉讼请求。')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('此致')
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('[XX市XX区]人民法院')
p.alignment = WD_ALIGN_PARAGRAPH.LEFT

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('具状人（原告签字）：')
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

p = doc.add_paragraph()
p.add_run('[XXXX年XX月XX日]')
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# 保存
output_path = '/Users/wangming/.openclaw/workspace/民事起诉状_电梯卡每月授权纠纷.docx'
doc.save(output_path)
print(f"DOCX generated: {output_path}")
