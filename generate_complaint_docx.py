#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# 标题
title = doc.add_heading('民事起诉状', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# 原被告信息
p = doc.add_paragraph()
p.add_run('原告：').bold = True
p.add_run('[你的姓名]，性别：[男/女]，[出生年月日]，民族：[XX]，身份证号：[XXXXXXXXXXXXXXXXXX]，住址：[XX省XX市XX区XX小区XX号楼XX单元XX室]，联系电话：[XXXXXXXXXXX]')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('被告：').bold = True
p.add_run('[物业公司全称]，统一社会信用代码：[XXXXXXXXXXXXXXXXXX]，住所地：[XX省XX市XX区XX小区XX号楼XX单元XX室]，法定代表人：[XXX]，职务：[总经理/董事长]，联系电话：[XXXXXXXXXXX]')

doc.add_paragraph()
doc.add_heading('诉讼请求', level=2)

doc.add_paragraph('1.  请求依法确认被告以未缴纳物业费为由禁止原告参与小区车位摇号的行为无效；')
doc.add_paragraph('2.  请求判令被告立即允许原告参与[小区名称]车位摇号分配；')
doc.add_paragraph('3.  请求判令被告赔偿原告经济损失人民币[XXXX]元（因无法摇号被迫以更高价格租赁车位产生的额外费用，若无损失可删除此项）；')
doc.add_paragraph('4.  本案诉讼费用由被告承担。')

doc.add_paragraph()
doc.add_heading('事实与理由', level=2)

# 一、基本案情
p = doc.add_paragraph()
p.add_run('一、基本案情').bold = True

doc.add_paragraph(
'\n原告系[XX市XX区XX小区]XX号楼XX单元XX室业主，依法对小区共有部分享有建筑物区分所有权。\n\n'
'[XXXX年XX月XX日]，被告发布《[小区名称]车位摇号分配通知》，明确规定\"未缴纳物业费的业主不得参与本次车位摇号分配\"，将缴纳物业费作为参与车位摇号的前置条件。\n\n'
'原告因[简述原因，如\"对物业服务存在争议暂缓缴费\"/\"暂时资金周转困难\"]未缴纳当期物业费，被告据此拒绝原告参与车位摇号，导致原告无法依法行使对小区共有车位的共有和管理权利。'
)

# 二、被告行为违法
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('二、被告行为违反法律规定，应认定无效').bold = True

doc.add_paragraph()
doc.add_paragraph('1. 两种法律关系相互独立，被告无权捆绑限制权利')
doc.add_paragraph('• 缴纳物业费是原告基于《物业服务合同》产生的合同义务；')
doc.add_paragraph('• 参与车位摇号是原告基于《中华人民共和国民法典》第二百七十一条、第二百七十五条规定，作为业主对小区共有车位享有的法定共有权；')
doc.add_paragraph('• 二者是完全独立的法律关系，被告无权将\"行使法定权利\"作为\"履行合同义务\"的前置条件，该捆绑行为缺乏法律依据。\n')

doc.add_paragraph('2. 小区车位属全体业主共有，被告越权限制原告权利')
doc.add_paragraph('案涉摇号分配的车位，系占用小区业主共有道路/场地设置，根据《民法典》第二百七十五条规定，属于业主共有。车位分配属于业主共同决定事项，被告未经业主大会同意，擅自以未缴费为由剥夺原告参与权，属于越权行为，严重侵犯原告的共有权。\n')

doc.add_paragraph('3. 违反公平自愿原则，构成不当限制')
doc.add_paragraph('被告该做法将物业费催收与业主基本权利挂钩，违反民事活动公平、自愿原则，对业主权利构成不合理限制。\n')

doc.add_paragraph('4. 违反法律关于物业费催收的禁止性规定')
doc.add_paragraph(
'《中华人民共和国民法典》第九百四十四条明确规定：\"业主逾期不支付物业费的，物业服务人可以催告其在合理期限内支付；合理期限届满仍不支付的，可以提起诉讼或者申请仲裁。禁止物业服务人采取停止供电、供水、供热、供燃气等方式催交物业费。\"\n\n'
'尽管法律未明文列举\"限制车位摇号\"，但该行为本质上仍是以限制业主法定权利的方式催交物业费，与法律禁止的停止供水供电性质相同，司法实践普遍认定此种行为无效。'
)

# 三、原告损失
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('三、原告损失').bold = True
doc.add_paragraph(
'因被告拒绝原告参与摇号，原告无法获得价格优惠的小区车位，被迫[描述具体损失，如\"以每月XXX元的价格在外租赁车位，比小区内部车位每月多支出XXX元\"]，截至起诉之日已产生额外损失共计人民币[XXXX]元，依法应由被告赔偿。'
)

doc.add_paragraph()
doc.add_heading('证据清单', level=2)

doc.add_paragraph('1.  业主身份证明：房产证/购房合同复印件——证明原告系小区业主，主体适格；')
doc.add_paragraph('2.  被告发布的车位摇号通知——证明被告明确将缴纳物业费作为参与摇号的前置条件；')
doc.add_paragraph('3.  沟通记录：微信/短信/录音——证明被告实际拒绝原告参与摇号；')
doc.add_paragraph('4.  [额外损失证据]：外部租赁合同、付款凭证——证明原告额外支出。')

doc.add_paragraph()
doc.add_heading('法律依据', level=2)
doc.add_paragraph('- 《中华人民共和国民法典》第二百七十一条、第二百七十五条、第九百四十四条；')
doc.add_paragraph('- 《物业管理条例》第六条；')
doc.add_paragraph('- 最高人民法院《关于审理物业服务纠纷案件适用法律若干问题的解释》相关规定。')

doc.add_paragraph()

doc.add_paragraph('综上所述，被告行为违反法律规定，侵犯原告合法权益，恳请贵院依法支持原告诉讼请求。')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('此致')
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('[XX市XX区]人民法院')
p.alignment = WD_ALIGN_PARAGRAPH.LEFT

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('具状人（原告签字）：')
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

p = doc.add_paragraph()
p.add_run('[XXXX年XX月XX日]')
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# 保存
output_path = '/Users/wangming/.openclaw/workspace/民事起诉状_车位摇号权利纠纷.docx'
doc.save(output_path)
print(f"DOCX generated: {output_path}")
