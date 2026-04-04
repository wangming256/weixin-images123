#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# 标题
title = doc.add_heading('民事起诉状', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# 原被告信息
p = doc.add_paragraph()
p.add_run('原告：').bold = True
p.add_run('[你的姓名]，性别：[男/女]，[出生年月日]，民族：[XX]，身份证号：[XXXXXXXXXXXXXXXXXX]，住址：[XX省XX市XX区XX小区XX号楼XX单元XX室]，联系电话：[XXXXXXXXXXX]')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('被告：').bold = True
p.add_run('[物业公司全称]，统一社会信用代码：[XXXXXXXXXXXXXXXXXX]，住所地：[XX省XX市XX区XX小区XX号楼XX单元XX室]，法定代表人：[XXX]，职务：[总经理/董事长]，联系电话：[XXXXXXXXXXX]')

doc.add_paragraph()
doc.add_heading('诉讼请求', level=2)

doc.add_paragraph('1.  请求依法确认被告对原告采取"未缴纳物业费必须每月到物业升级电梯卡，正常缴纳物业费可按年升级"的区别对待行为违法，并认定该行为无效；')
doc.add_paragraph('2.  请求判令被告立即给予原告与正常缴费业主同等的电梯卡授权期限（至少一年一升），不得对原告区别对待；')
doc.add_paragraph('3.  请求判令被告赔偿原告因每月必须前往物业升级产生的交通费、误工费共计人民币[XXXX]元；')
doc.add_paragraph('4.  本案诉讼费用由被告承担。')

doc.add_paragraph()
doc.add_heading('事实与理由', level=2)

# 一、基本案情
p = doc.add_paragraph()
p.add_run('一、基本案情').bold = True

doc.add_paragraph(
'\n原告系[XX市XX区XX小区]XX号楼XX单元XX室业主，依法对房屋享有所有权，并有权正常使用小区公共设施包括电梯。\n\n'
'该小区电梯卡授权升级规则：**正常缴纳物业费的业主，可以一次性授权一年，每年只需要去物业升级一次；但对于未全额缴纳当期物业费的业主，被告要求必须**每月都到物业办公地点手动升级一次**，不升级到期电梯卡就不能使用电梯。\n\n'
'原告[简述情况：因故暂缓缴纳本期物业费/因物业服务争议暂未全额缴纳]，被告据此对原告执行区别对待，要求原告必须每月亲自跑一趟物业才能升级，否则到期就无法使用电梯。'
'原告认为，被告这种区别对待、人为制造麻烦逼迫缴费的做法，违反法律规定，侵犯原告合法权益，故诉至法院。'
)

# 二、被告行为违法
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('二、被告行为违反法律规定，构成侵权').bold = True

doc.add_paragraph()
doc.add_paragraph('1. 使用电梯是法定权利，被告无权对欠缴物业费业主区别对待、增加负担')
doc.add_paragraph(
'• 根据《民法典》第二百七十一条，业主对小区共有部分享有共有和共同管理的权利，电梯属于共有公共设施，业主依法有权正常使用。\n'
'• 即便业主欠缴物业费，被告也应当保障业主正常使用公共设施，其有权通过诉讼途径追讨物业费，无权在使用电梯问题上对业主区别对待、人为制造不便逼迫缴费。\n'
'• 正常缴费一年升级一次，未缴费就要每月跑一趟，这种差别对待本身就是对业主权利的不当限制，没有任何法律依据。'
)

doc.add_paragraph()
doc.add_paragraph('2. 这种行为本质是以限制使用方式变相催费，违反民法典禁止性规定')
doc.add_paragraph(
'《中华人民共和国民法典》第九百四十四条第三款明确规定：\n\n'
'> **业主逾期不支付物业费的，物业服务人可以催告其在合理期限内支付；合理期限届满仍不支付，可以提起诉讼或者申请仲裁。**\n'
'> **物业服务人不得采取停止供电、供水、供热、供燃气等方式催交物业费。**\n\n'
'被告这种"每月必须亲自升级，不升级就不能用电梯"的做法，本质就是通过给业主日常生活制造不便来逼迫业主缴费，'
'属于法律禁止的"以限制业主使用公共设施的方式催交物业费"，性质上和法律明文禁止的停水停电没有区别，应当认定无效。'
)

doc.add_paragraph()
doc.add_paragraph('3. 该行为已经给原告造成实际损失')
doc.add_paragraph(
'原告每月必须专程请假前往物业办理升级，已经实际产生[简述："误工费每月XXX元、交通费XXX元，共计XXX元"]，该损失完全是被告违法行为造成，依法应由被告赔偿。'
)

doc.add_paragraph()
doc.add_heading('证据清单', level=2)

doc.add_paragraph('1.  业主身份证明：房产证/购房合同/不动产权证——证明原告系小区业主，主体适格；')
doc.add_paragraph('2.  被告电梯卡规则证据：被告公示文件/与物业工作人员沟通记录/其他业主证明——证明被告对正常缴费业主和欠缴业主采取不同升级周期，区别对待；')
doc.add_paragraph('3.  原告损失证据：交通费票据/误工证明——证明原告实际损失金额；')
doc.add_paragraph('4.  其他：[如有其他证据补充]。')

doc.add_paragraph()
doc.add_heading('法律依据', level=2)
doc.add_paragraph('- 《中华人民共和国民法典》第二百七十一条、第九百四十四条、第一千一百六十五条；')
doc.add_paragraph('- 《物业管理条例》第六条。')

doc.add_paragraph()

doc.add_paragraph('综上所述，被告对欠缴物业费业主采取每月强制升级的区别对待行为，本质是变相限制业主使用电梯催缴物业费，违反民法典明确禁止性规定，侵害原告合法权益，恳请贵院依法支持原告诉讼请求。')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('此致')
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('[XX市XX区]人民法院')
p.alignment = WD_ALIGN_PARAGRAPH.LEFT

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('具状人（原告签字）：')
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

p = doc.add_paragraph()
p.add_run('[XXXX年XX月XX日]')
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# 保存
output_path = '/Users/wangming/.openclaw/workspace/民事起诉状_电梯卡区别对待纠纷.docx'
doc.save(output_path)
print(f"DOCX generated: {output_path}")
