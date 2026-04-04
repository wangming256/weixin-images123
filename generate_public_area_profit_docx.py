#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# 标题
title = doc.add_heading('民事起诉状', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# 原被告信息
p = doc.add_paragraph()
p.add_run('原告：').bold = True
p.add_run('[你的姓名]，性别：[男/女]，[出生年月日]，民族：[XX]，身份证号：[XXXXXXXXXXXXXXXXXX]，住址：[XX省XX市XX区XX小区XX号楼XX单元XX室]，联系电话：[XXXXXXXXXXX]')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('被告：').bold = True
p.add_run('[物业公司全称]，统一社会信用代码：[XXXXXXXXXXXXXXXXXX]，住所地：[XX省XX市XX区XX小区XX号楼XX单元XX室]，法定代表人：[XXX]，职务：[总经理/董事长]，联系电话：[XXXXXXXXXXX]')

doc.add_paragraph()
doc.add_heading('诉讼请求', level=2)

doc.add_paragraph('1.  请求判令被告立即向原告公开[XXXX年XX月至XXXX年XX月]小区公共区域收益明细（包括但不限于公共区域广告收入、车位租金、场地租赁费、公共设施使用费等）；')
doc.add_paragraph('2.  请求判令本案诉讼费用由被告承担。')

doc.add_paragraph()
doc.add_heading('事实与理由', level=2)

# 一、基本案情
p = doc.add_paragraph()
p.add_run('一、基本案情').bold = True

doc.add_paragraph(
'\n原告系[XX市XX区XX小区]XX号楼XX单元XX室业主，依法对小区公共区域享有共有权和知情权。\n\n'
'被告系该小区物业服务企业，根据《物业管理条例》《民法典》相关规定，小区公共区域产生的收益属于全体业主共有，'
'物业服务企业应当定期向业主公开公共收益的收支情况，接受业主监督。\n\n'
'原告[XXXX年XX月XX日]已通过[书面/微信/短信]方式向被告提出申请，要求公开[XXXX年XX月至XXXX年XX月]期间小区公共区域收益明细，'
'但被告至今拒绝公开，也未提供任何明细信息，原告认为被告行为侵犯原告知情权，故诉至法院。'
)

# 二、被告行为违法
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('二、被告行为违反法律规定，侵害业主知情权').bold = True

doc.add_paragraph()
doc.add_paragraph('1. 公共区域收益属于全体业主共有，业主依法享有知情权')
doc.add_paragraph(
'• 根据《中华人民共和国民法典》第二百七十一条、第二百八十二条，建设单位、物业服务企业或者其他管理人等利用业主的共有部分产生的收入，'
'在扣除合理成本之后，属于业主共有。\n'
'• 根据《物业管理条例》第六条，业主享有"对物业共用部位、共用设施设备和相关场地使用情况享有知情权和监督权"。\n'
'• 业主知情权是法定权利，被告作为管理人，有义务公开公共收益明细，接受业主监督。'
)

doc.add_paragraph()
doc.add_paragraph('2. 原告已经履行申请程序，被告拒绝公开构成侵权')
doc.add_paragraph(
'原告已经依法向被告提出公开申请，被告无正当理由拒绝公开，拒不履行法定义务，'
'已经实际侵害原告对共有部分的知情权，原告有权通过诉讼途径要求被告履行公开义务。'
)

doc.add_paragraph()
doc.add_heading('证据清单', level=2)

doc.add_paragraph('1.  业主身份证明：房产证/购房合同/不动产权证——证明原告系小区业主，主体适格；')
doc.add_paragraph('2.  原告向被告申请公开的证据：申请书/沟通记录（微信/短信/邮件/书面签收记录）——证明原告已经提出申请，被告拒绝公开；')
doc.add_paragraph('3.  其他：[如有其他证据补充]。')

doc.add_paragraph()
doc.add_heading('法律依据', level=2)
doc.add_paragraph('- 《中华人民共和国民法典》第二百七十一条、第二百八十二条；')
doc.add_paragraph('- 《物业管理条例》第六条、第五十四条；')
doc.add_paragraph('- 《最高人民法院关于审理建筑物区分所有权纠纷案件适用法律若干问题的解释》第十三条。')

doc.add_paragraph()

doc.add_paragraph('综上所述，被告作为小区物业服务企业，拒不向业主公开公共区域收益明细，侵害业主法定知情权，恳请贵院依法支持原告诉讼请求。')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('此致')
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('[XX市XX区]人民法院')
p.alignment = WD_ALIGN_PARAGRAPH.LEFT

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('具状人（原告签字）：')
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

p = doc.add_paragraph()
p.add_run('[XXXX年XX月XX日]')
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# 保存
output_path = '/Users/wangming/.openclaw/workspace/民事起诉状_公共区域收益知情权纠纷.docx'
doc.save(output_path)
print(f"DOCX generated: {output_path}")
