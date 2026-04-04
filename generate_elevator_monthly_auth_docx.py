#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# 标题
title = doc.add_heading('民事起诉状', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# 原被告信息
p = doc.add_paragraph()
p.add_run('原告：').bold = True
p.add_run('[你的姓名]，性别：[男/女]，[出生年月日]，民族：[XX]，身份证号：[XXXXXXXXXXXXXXXXXX]，住址：[XX省XX市XX区XX小区XX号楼XX单元XX室]，联系电话：[XXXXXXXXXXX]')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('被告：').bold = True
p.add_run('[物业公司全称]，统一社会信用代码：[XXXXXXXXXXXXXXXXXX]，住所地：[XX省XX市XX区XX小区XX号楼XX单元XX室]，法定代表人：[XXX]，职务：[总经理/董事长]，联系电话：[XXXXXXXXXXX]')

doc.add_paragraph()
doc.add_heading('诉讼请求', level=2)

doc.add_paragraph('1.  请求依法确认被告故意将电梯卡设置为"必须每月到物业升级授权"，并以未缴纳物业费为由拒绝升级，导致原告无法正常使用电梯的行为违法无效；')
doc.add_paragraph('2.  请求判令被告立即为原告电梯卡办理长期/永久授权，不得再以每月升级为由限制原告使用电梯；')
doc.add_paragraph('3.  请求判令被告赔偿原告经济损失人民币[XXXX]元（因每月必须跑腿产生的交通费、误工费等）；')
doc.add_paragraph('4.  本案诉讼费用由被告承担。')

doc.add_paragraph()
doc.add_heading('事实与理由', level=2)

# 一、基本案情
p = doc.add_paragraph()
p.add_run('一、基本案情').bold = True

doc.add_paragraph(
'\n原告系[XX市XX区XX小区]XX号楼XX单元XX室业主，依法对房屋享有所有权，并有权正常使用小区公共设施包括电梯。\n\n'
'案涉小区电梯卡本可以设置长期/永久授权，原告入住后，被告却**故意设置为必须每月到被告办公地点手动升级授权一次**——'
'如果原告当月物业费未缴，被告就拒绝给电梯卡升级，电梯卡到期后即无法使用电梯。\n\n'
'原告[简述情况：因故暂缓缴纳本期物业费/因物业服务争议暂未缴纳]，被告据此拒绝为原告电梯卡进行本月授权升级，导致原告电梯卡失效无法正常使用电梯。'
'原告认为，被告这种故意设置 monthly 授权门槛绑定物业费催收的行为，已经严重侵犯原告合法权益，故诉至法院。'
)

# 二、被告行为违法
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('二、被告行为违反法律规定，构成侵权').bold = True

doc.add_paragraph()
doc.add_paragraph('1. 电梯是公共设施，业主有权无偿正常使用，不应被附加物业费绑定条件')
doc.add_paragraph(
'• 根据《民法典》第二百七十一条，业主对建筑物内的住宅等专有部分享有所有权，对专有部分以外的共有部分享有共有和共同管理的权利。\n'
'• 电梯属于小区共有公共设施，业主正常使用电梯是法定权利，该项权利与业主是否缴纳物业费没有法律上的牵连关系。\n'
'• 被告作为物业服务企业，本身就负有保障小区公共设施正常使用的义务，无权故意给业主正常使用电梯设置人为门槛，更无权以\"不缴费就不给升级\"来胁迫业主缴费。'
)

doc.add_paragraph()
doc.add_paragraph('2. 这种行为本质就是以限制使用电梯的方式催交物业费，违反民法典明确禁止性规定')
doc.add_paragraph(
'《中华人民共和国民法典》第九百四十四条第三款明确规定：\n\n'
'> **业主逾期不支付物业费的，物业服务人可以催告其在合理期限内支付；合理期限届满仍不支付的，可以提起诉讼或者申请仲裁。**\n'
'> **物业服务人不得采取停止供电、供水、供热、供燃气等方式催交物业费。**\n\n'
'被告明知法律禁止停止供应必需公共服务催交物业费，却采用"故意设置每月授权门槛"这种"软"方式——'
'名义上没有"停止电梯"，实际上就是"不缴费就每月让你用不了"，本质和直接停运电梯没有区别，'
'这是**变相违反法律禁止性规定**，行为应当认定无效。'
)

doc.add_paragraph()
doc.add_paragraph('3. 被告该行为给原告造成额外负担和损失')
doc.add_paragraph(
'正常情况下，业主拿到电梯卡后可以一直使用，不需要每月跑物业一趟。被告这种设计完全是人为给业主增加麻烦：'
'\n\n• 每月必须专程请假跑腿去物业，产生误工费、交通费；\n• 一旦忘记或者暂时没钱缴费，直接无法使用电梯，严重影响日常生活；\n• 本质就是通过给业主制造不便来逼迫缴费，违背物业服务企业基本义务。'
)

# 三、损失
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('三、原告损失').bold = True
doc.add_paragraph(
'因被告拒绝升级，原告[简述具体损失："每月产生交通费XX元、误工费XX元，截至起诉共计XX元"/"因无法使用电梯不得不打车出行，产生额外费用XX元"]，'
'该损失系被告违法行为造成，依法应由被告赔偿。'
)

doc.add_paragraph()
doc.add_heading('证据清单', level=2)

doc.add_paragraph('1.  业主身份证明：房产证/购房合同/不动产权证——证明原告系小区业主，主体适格；')
doc.add_paragraph('2.  电梯卡必须每月升级、被告拒绝升级的证据：小区电梯卡规则/被告公示/与被告沟通记录（微信/短信/录音/录像）——证明被告设置每月升级并以此拒绝原告使用电梯；')
doc.add_paragraph('3.  [损失证据：交通费票据、误工证明等——证明原告实际损失]；')
doc.add_paragraph('4.  其他：[如有其他证据补充]。')

doc.add_paragraph()
doc.add_heading('法律依据', level=2)
doc.add_paragraph('- 《中华人民共和国民法典》第二百七十一条、第九百四十四条、第一千一百六十五条；')
doc.add_paragraph('- 《物业管理条例》第六条。')

doc.add_paragraph()

doc.add_paragraph('综上所述，被告故意将电梯卡设置为每月必须到物业升级，并以此绑定物业费催收，变相限制原告正常使用电梯，违反民法典明确禁止性规定，严重侵害原告合法权益，恳请贵院依法支持原告诉讼请求。')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('此致')
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('[XX市XX区]人民法院')
p.alignment = WD_ALIGN_PARAGRAPH.LEFT

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('具状人（原告签字）：')
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

p = doc.add_paragraph()
p.add_run('[XXXX年XX月XX日]')
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# 保存
output_path = '/Users/wangming/.openclaw/workspace/民事起诉状_电梯卡每月故意设置授权升级纠纷.docx'
doc.save(output_path)
print(f"DOCX generated: {output_path}")
