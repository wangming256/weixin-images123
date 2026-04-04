#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# 标题
title = doc.add_heading('民事起诉状', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# 原被告信息
p = doc.add_paragraph()
p.add_run('原告：').bold = True
p.add_run('[你的姓名]，性别：[男/女]，[出生年月日]，民族：[XX]，身份证号：[XXXXXXXXXXXXXXXXXX]，住址：[沈阳市沈河区金家街98号锦城邻里XX号楼XX单元XX室]，联系电话：[13700008144]')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('被告：').bold = True
p.add_run('[辽宁中恒物业管理有限公司]，统一社会信用代码：[XXXXXXXXXXXXXXXXXX]，住所地：[沈阳市沈河区金家街98号锦城邻里]，法定代表人：[XXX]，职务：[总经理/董事长]，联系电话：[XXXXXXXXXXX]')

doc.add_paragraph()
doc.add_heading('诉讼请求', level=2)

doc.add_paragraph('1.  请求判令被告向原告公开公示下列信息：')
doc.add_paragraph('   (1) 营业执照、建筑区划总平面图、管理人员持证上岗证书、小区管理规约、服务人员的基本情况、职责分工、联系方式、投诉机制、物业服务合同、物业服务等级、服务内容、服务标准、收费项目、收费依据、收费方式、计费起始时间、价格举报电话、消防等设备设施安全生产应急预案等物业管理工作必须公开事项；')
doc.add_paragraph('   (2)  2021年7月至今，被告利用小区公共区域、共用部位、公用设施设备的经营收益明细；')
doc.add_paragraph('   (3)  2021年7月至今，公共区域、共用部位养护记录，建筑物及其附属设施的维修资金筹集使用情况；')
doc.add_paragraph('   (4)  小区停车管理方案、现有车位使用人名单/出租明细；')
doc.add_paragraph('   (5)  公共区域绿化养护记录、楼体外墙/公共区域/楼顶防水养护记录；')
doc.add_paragraph('   (6)  物业服务用房使用情况；')
doc.add_paragraph('')
doc.add_paragraph('2.  请求判令被告退还重复收取原告的电梯维护/年检分摊费用人民币[XXXX]元；')
doc.add_paragraph('')
doc.add_paragraph('3.  请求判令本案诉讼费用由被告承担。')

doc.add_paragraph()
doc.add_heading('事实与理由', level=2)

# 一、基本案情
p = doc.add_paragraph()
p.add_run('一、基本案情').bold = True

doc.add_paragraph(
'\n原告系沈阳市沈河区锦城邻里小区[XX号楼XX单元XX室]业主，依法对小区公共部分享有建筑物区分所有权和知情权。被告系该小区物业服务企业，双方签订有书面物业服务合同。\n\n'
'原告为行使知情权，已于[2024年9月5日]向被告发出书面申请函，要求被告公开前述诉讼请求列明的全部信息，但被告至今未完整公开，'
'且原告在被告公开的部分年度收支公示中发现被告存在重复列支电梯费用、车位管理费收入与实际车位数量严重不符等问题，具体如下：'
)

# 二、被告侵犯原告知情权，拒不公开信息
p = doc.add_paragraph()
p.add_run('二、被告拒绝完整公开信息，侵犯原告法定知情权').bold = True

doc.add_paragraph(
'\n根据《中华人民共和国民法典》第二百七十一条、第二百八十二条，《最高人民法院关于审理建筑物区分所有权纠纷案件适用法律若干问题的解释》第十三条，'
'业主对上述信息依法享有知情权和查阅权，原告已经依法向被告提出书面申请，被告无正当理由拒绝提供，已经构成侵权。'
)

# 三、被告存在重复收费，应当退还
p = doc.add_paragraph()
p.add_run('三、被告重复列支电梯费用，构成重复收费，应当退还').bold = True

doc.add_paragraph(
'\n根据双方签订的物业服务合同约定，原告交纳的物业费[1.9元/月/平方米]已经包含电梯费用。'
'但被告一方面在物业费支出中列支电梯维护费用[2022年一年即列支131,010元]，另一方面又单独向业主分摊收取电梯年检/维护费用，'
'属于就同一服务重复收费，原告已交纳的分摊费用应当退还。'
)

# 四、车位管理费收入账实不符，侵犯业主权益
p = doc.add_paragraph()
p.add_run('四、车位管理费收入与公示不符，原告有权知情').bold = True

doc.add_paragraph(
'\n根据被告公示，2022年车位管理费收入共计193,680元，按照合同约定每个车位每月120元计算，'
'该金额对应约134个车位，但小区实际共有车位251个，一年应收车位管理费应为361,440元，缺口高达约168,000元，'
'被告至今无法提供完整出租明细，原告作为业主无法知晓真实收支情况，有权要求公开。'
)

doc.add_paragraph()
doc.add_heading('证据清单', level=2)

doc.add_paragraph('1.  原告房产证——证明原告系小区业主，主体适格；')
doc.add_paragraph('2.  物业服务合同——证明原被告存在物业服务合同关系，物业费含电梯费、公共收益分配约定；')
doc.add_paragraph('3.  原告向被告发出的信息公开申请函——证明原告已经依法申请公开，被告未予完整公开；')
doc.add_paragraph('4.  原告与被告工作人员微信聊天记录——证明原告多次索要信息，被告拒绝提供；')
doc.add_paragraph('5.  被告公示的年度物业服务收支、公共收益公示——证明被告重复列支电梯费、车位收入缺口；')
doc.add_paragraph('6.  原告交纳电梯分摊费用的凭证——证明原告实际交纳了该笔费用，被告重复收费。')

doc.add_paragraph()
doc.add_heading('法律依据', level=2)

doc.add_paragraph('- 《中华人民共和国民法典》第二百七十一条、第二百八十二条；')
doc.add_paragraph('- 《物业管理条例》第六条；')
doc.add_paragraph('- 《最高人民法院关于审理建筑物区分所有权纠纷案件适用法律若干问题的解释》第十三条。')

doc.add_paragraph()

doc.add_paragraph(
'综上所述，被告拒绝向原告公开物业管理相关信息，违反法律规定和合同约定，'
'且存在重复收费行为，侵犯原告合法权益，原告依法向贵院提起诉讼，恳请贵院依法支持原告诉讼请求。'
)

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('此致')
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('[沈阳市沈河区]人民法院')
p.alignment = WD_ALIGN_PARAGRAPH.LEFT

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('具状人（原告签字）：')
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

p = doc.add_paragraph()
p.add_run('[XXXX年XX月XX日]')
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# 保存
output_path = '/Users/wangming/.openclaw/workspace/民事起诉状_业主知情权纠纷.docx'
doc.save(output_path)
print(f"DOCX generated: {output_path}")
