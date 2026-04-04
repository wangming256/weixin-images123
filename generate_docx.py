#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_ALIGN_VERTICAL

doc = Document()

# 标题
title = doc.add_heading('业主大会议事规则（草案）合规性问题对照表', level=1)

# 基本信息
doc.add_paragraph()
doc.add_paragraph('文件来源：业主大会议事规则（草案）')
doc.add_paragraph('对照标准：住建部《业主大会和业主委员会指导规则》（建房[2009]274号）')
doc.add_paragraph('生成日期：2026年4月1日')
doc.add_paragraph()

# 一、整体评价
doc.add_heading('一、整体评价', level=2)
table = doc.add_table(rows=2, cols=2)
table.cell(0, 0).text = '项目'
table.cell(0, 1).text = '结论'
table.cell(1, 0).text = '整体框架'
table.cell(1, 1).text = '✅ 符合要求'
table.cell(0, 0).paragraphs[0].runs[0].font.bold = True
table.cell(0, 1).paragraphs[0].runs[0].font.bold = True
doc.add_paragraph()

# 二、问题清单
doc.add_heading('二、问题清单', level=2)
table = doc.add_table(rows=1, cols=4)
table.autofit = True
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '序号'
hdr_cells[1].text = '问题描述'
hdr_cells[2].text = '国家规定要求'
hdr_cells[3].text = '修正建议'
for cell in hdr_cells:
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True

problems = [
    ('1', '未明确业主委员会人数', '业主委员会由 5-11 人组成，且必须是单数', '补充明确具体人数（例：本小区业主委员会由 7 名委员组成）'),
    ('2', '换届时间错误', '业主委员会任期届满前 3 个月 应当进行换届选举', '草案写的是「任期届满两个月前」，应修正为 三个月'),
    ('3', '缺失投票权确定规则', '已竣工但尚未出售或者尚未交给物业买受人的物业，建设单位享有一票投票权', '补充本条规定'),
    ('4', '缺失居委会告知程序', '业主大会会议应当 书面告知居民委员会', '补充：会议通知同时抄送居民委员会'),
    ('5', '缺失公告期限要求', '业主大会决定作出后，公告期 不少于 7 日', '补充公告期限要求'),
    ('6', '缺失逾期换届处理机制', '业主委员会逾期不换届，可由 物业所在地的区、县房地产行政主管部门或者街道办事处、乡镇人民政府指导业主重新选举', '补充本条处理机制'),
    ('7', '（可选）缺失候补委员制度', '国家提倡建立业主委员会候补委员制度', '可补充：设立候补委员不超过 5 人，缺额依次递补'),
]

for p in problems:
    row_cells = table.add_row().cells
    row_cells[0].text = p[0]
    row_cells[1].text = p[1]
    row_cells[2].text = p[2]
    row_cells[3].text = p[3]
    for cell in row_cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

doc.add_paragraph()

# 三、符合规定的条款说明
doc.add_heading('三、符合规定的条款说明', level=2)
doc.add_paragraph('以下内容原草案已经符合国家规定：')
items = [
    '1. ✅ 业主委员会任期 3 年 → 符合「不超过 5 年」要求',
    '2. ✅ 业主大会会议召开条件 → 符合「必须有 1/2 以上投票表决权业主参加方为有效」',
    '3. ✅ 表决通过比例 → 普通决议 1/2、特别决议 2/3，完全符合规定',
    '4. ✅ 临时会议召开情形 → 「20% 以上业主提议应当召开」符合规定',
    '5. ✅ 经费公开制度 → 「定期公告、接受业主质询」符合规定',
    '6. ✅ 印章管理制度 → 「登记备案、使用登记」符合规定',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()

# 四、结论
doc.add_heading('四、结论', level=2)
doc.add_paragraph('原草案主体内容框架正确，核心表决规则符合要求，仅缺少若干程序性细节，补充修正后即可完全符合《业主大会和业主委员会指导规则》要求。')

doc.add_paragraph()
doc.add_paragraph('生成：OpenClaw AI 助手')
doc.add_paragraph('OCR 识别：完成 7 页全部文字')
doc.add_paragraph('合规性检查：对照国家规定完成')

# 保存
doc.save('/Users/wangming/.openclaw/workspace/业主大会议事规则_问题对照表.docx')
print('DOCX 生成成功')
